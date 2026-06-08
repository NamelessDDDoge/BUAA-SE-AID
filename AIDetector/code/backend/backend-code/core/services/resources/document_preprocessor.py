from pathlib import Path
import re

from .text_sanitizer import sanitize_text_content


UNREADABLE_FILE_MESSAGE = "无法读取文件内容，请上传可解析的文本文件。"
EMPTY_FILE_MESSAGE = "无内容"


def preprocess_document(file_path, max_segment_length=1500, overlap_length=100, fallback_segment_length=40000):
    tables = extract_document_tables(file_path)
    text_content = sanitize_text_content(extract_document_text(file_path))
    paragraphs = extract_document_paragraphs(text_content)
    
    sections = parse_document_sections(text_content)
    # 将所有非参考文献的部分拼接作为待检测的正文内容
    core_text = sections.get("abstract", "") + "\n\n" + sections.get("body", "") + "\n\n" + sections.get("acknowledgements", "")
    core_text = core_text.strip()
    
    if not core_text:
        core_text = text_content
        
    # 应用学术噪音清洗 (去除图表标题、公式乱码、文内引用)
    core_text = sanitize_academic_noise(core_text)
        
    segments = split_text_into_segments(
        core_text,
        max_segment_length=max_segment_length,
        overlap_length=overlap_length,
        fallback_segment_length=fallback_segment_length,
    )
    
    # 调试日志：按照你的要求，将分段后的结果和提取原文直接输出到服务器本地文件中
    try:
        debug_file = str(file_path) + ".debug_segments.txt"
        with open(debug_file, "w", encoding="utf-8") as f:
            f.write("=== 【1】从 Word/PDF 提取到的原始文本 (前2000字) ===\n")
            f.write(text_content[:2000] + "\n\n...\n\n")
            f.write("=== 【2】文档区块(Section)解析结果 ===\n")
            f.write(f"摘要 (abstract) 解析出: {len(sections.get('abstract', '').split(chr(10)))} 行\n")
            f.write(f"正文 (body) 解析出: {len(sections.get('body', '').split(chr(10)))} 行\n")
            f.write(f"致谢 (acknowledgements) 解析出: {len(sections.get('acknowledgements', '').split(chr(10)))} 行\n")
            f.write(f"参考文献 (references) 解析出: {len(sections.get('references', '').split(chr(10)))} 行\n\n")
            f.write("=== 【3】最终送去检测的段落切分结果 ===\n")
            for i, seg in enumerate(segments):
                f.write(f"--- 第 {i+1} 段 ---\n{seg}\n\n")
            f.write("=== 【4】参考文献切分结果 ===\n")
            refs = extract_document_references(text_content)
            for i, ref in enumerate(refs):
                f.write(f"--- 参考文献 {i+1} ---\n{ref}\n\n")
    except Exception:
        pass

    return {
        "text_content": text_content,
        "paragraphs": paragraphs,
        "sections": sections,
        "references": extract_document_references(text_content),
        "segments": segments,
        "tables": tables,
    }


def extract_document_text(file_path):
    file_ext = Path(file_path).suffix.lower()
    try:
        if file_ext == ".pdf":
            import fitz

            text_parts = []
            with fitz.open(file_path) as document:
                for page in document:
                    page_height = page.rect.height
                    table_regions = _get_pdf_table_regions(page)
                    blocks = page.get_text("blocks")
                    for b in blocks:
                        if b[6] == 0:  # text block
                            x0, y0, x1, y1, raw_text, block_no, block_type = b
                            if _intersects_any_bbox((x0, y0, x1, y1), table_regions):
                                continue
                            raw_text = raw_text.strip()
                            if not raw_text:
                                continue
                                
                            # 剔除页眉、页脚和页码 (基于页面坐标高度系)
                            is_header = y1 < page_height * 0.08  # 整个文本块位于页面顶部 8% 区域
                            is_footer = y0 > page_height * 0.92  # 整个文本块位于页面底部 8% 区域
                            # 针对纯数字页码，稍微放宽一点区域到上下 12%
                            is_page_num = raw_text.isdigit() and (y1 < page_height * 0.12 or y0 > page_height * 0.88)
                            
                            if is_header or is_footer or is_page_num:
                                continue
                                
                            lines = raw_text.split('\n')
                            merged_lines = []
                            for line in lines:
                                line = line.strip()
                                if not line:
                                    continue
                                if not merged_lines:
                                    merged_lines.append(line)
                                else:
                                    prev_line = merged_lines[-1]
                                    # 处理连字符折行
                                    if prev_line.endswith('-'):
                                        merged_lines[-1] = prev_line[:-1] + line
                                    # 如果是列表/参考文献开头，则保留换行
                                    elif re.match(r'^(\[\d+\]|\(\d+\)|\d+\.)', line):
                                        # 但如果它仅仅只是一个孤立的数字（比如论文里单独占一行的年份 2023. ），它不是新段落
                                        if re.match(r'^(\[\d+\]|\(\d+\)|\d+\.)$', line) or (line.isdigit() or line.strip('.').isdigit()):
                                            merged_lines[-1] += " " + line
                                        else:
                                            merged_lines.append(line)
                                    # 如果当前行完全是数字，或者只有数字和标点符号（例如年份、页码），肯定是上一行折断下来的
                                    elif re.match(r'^[\d\s\.,\-]+$', line):
                                        merged_lines[-1] += " " + line
                                    # 否则认为是同一段落内部的折行，用空格拼合
                                    else:
                                        merged_lines[-1] += " " + line
                                        
                            # 跨 Block/跨页的段落拼合逻辑
                            for m_line in merged_lines:
                                if not text_parts:
                                    text_parts.append(m_line)
                                else:
                                    prev_part = text_parts[-1]
                                    # 如果上一块以连字符结尾
                                    if prev_part.endswith('-'):
                                        text_parts[-1] = prev_part[:-1] + m_line
                                    # 如果当前块完全是数字和标点，或者是孤立的数字，拼到上一块去
                                    elif re.match(r'^[\d\s\.,\-]+$', m_line) or (m_line.strip('.').isdigit()):
                                        text_parts[-1] = prev_part + " " + m_line
                                    else:
                                        # 检查上一块是否以正常的段落结束符（句号、问号、叹号、引号、冒号等）结尾
                                        ends_with_punct = re.search(r'[.!?。！？\”"\]\)：:]$', prev_part[-1:])
                                        # 如果没以结束符结尾，并且 (下一块是小写/标点开头，或者上一块字数较多说明是未写完的正文)
                                        if not ends_with_punct and (m_line[0].islower() or m_line[0] in ',;:' or len(prev_part) > 60):
                                            text_parts[-1] = prev_part + " " + m_line
                                        else:
                                            text_parts.append(m_line)
            return "\n".join(text_parts)
        if file_ext == ".docx":
            import docx

            document = docx.Document(file_path)
            # 原生的 document.paragraphs 无法读取文本框(Text Box)和表格内的文本。
            # PDF转Word的文档大量使用文本框排版，会导致正文大面积丢失。
            # 这里改用底层的 XPath 提取所有的 <w:p> (段落) 节点，确保无遗漏。
            # 过滤策略：只提取主文档（正文、文本框、表格）中的 <w:p>。
            # Word 的 XML 结构中，页眉和页脚被存放在独立的关系文件（header.xml, footer.xml）里，
            # 而 document.element (即 document.xml) 本身就不包含页眉页脚！
            # 因此，直接在 document.element 下执行 XPath，天然就免疫了页眉页脚和通过域代码生成的页码。
            texts = []
            for p in document.element.xpath('//w:p'):
                if _docx_paragraph_is_inside_table(p):
                    continue
                p_text = "".join(t.text for t in p.xpath('.//w:t') if t.text)
                if p_text.strip():
                    texts.append(p_text.strip())
            return "\n".join(texts)
        with open(file_path, "r", encoding="utf-8") as handle:
            return handle.read()
    except Exception:
        try:
            with open(file_path, "r", encoding="gbk") as handle:
                return handle.read()
        except Exception:
            return UNREADABLE_FILE_MESSAGE


def extract_document_tables(file_path):
    file_ext = Path(file_path).suffix.lower()
    try:
        if file_ext == ".pdf":
            return extract_pdf_tables(file_path)
        if file_ext == ".docx":
            return extract_docx_tables(file_path)
        if file_ext in {".txt", ".csv", ".tsv"}:
            return extract_text_tables(file_path)
    except Exception:
        return []
    return []


def extract_pdf_tables(file_path):
    import fitz

    tables = []
    with fitz.open(file_path) as document:
        for page_index, page in enumerate(document):
            native_regions = []
            for page_table_index, table in enumerate(_find_pdf_tables(page)):
                rows = _extract_pdf_table_rows(table)
                if rows:
                    bbox = getattr(table, "bbox", None)
                    if bbox and len(bbox) == 4:
                        native_regions.append(tuple(float(value) for value in bbox))
                    tables.append(
                        _build_table_payload(
                            len(tables),
                            rows,
                            source="pdf_native",
                            page_number=page_index + 1,
                            page_table_index=page_table_index,
                        )
                    )
            for inferred_index, inferred in enumerate(_infer_pdf_tables_from_words(page, exclude_bboxes=native_regions)):
                inferred["table_index"] = len(tables)
                inferred["page_number"] = page_index + 1
                inferred["page_table_index"] = inferred_index
                tables.append(inferred)
    return tables


def extract_docx_tables(file_path):
    import docx

    document = docx.Document(file_path)
    tables = []
    for table in document.tables:
        rows = []
        for row in table.rows:
            cells = [_normalize_table_cell(cell.text) for cell in row.cells]
            if any(cells):
                rows.append(cells)
        if rows:
            tables.append(_build_table_payload(len(tables), rows, source="docx"))
    return tables


def extract_text_tables(file_path):
    try:
        raw_text = Path(file_path).read_text(encoding="utf-8")
    except UnicodeDecodeError:
        raw_text = Path(file_path).read_text(encoding="gbk", errors="ignore")

    tables = []
    current_rows = []
    for line in raw_text.splitlines():
        parsed = _parse_text_table_row(line)
        if parsed:
            current_rows.append(parsed)
            continue
        if len(current_rows) >= 2:
            tables.append(_build_table_payload(len(tables), current_rows, source="text"))
        current_rows = []

    if len(current_rows) >= 2:
        tables.append(_build_table_payload(len(tables), current_rows, source="text"))
    return tables


def _find_pdf_tables(page):
    finder = getattr(page, "find_tables", None)
    if not callable(finder):
        return []
    try:
        result = finder()
    except Exception:
        return []
    return list(getattr(result, "tables", None) or [])


def _extract_pdf_table_rows(table):
    extractor = getattr(table, "extract", None)
    if not callable(extractor):
        return []
    try:
        rows = extractor()
    except Exception:
        return []
    return [
        [_normalize_table_cell(cell) for cell in row]
        for row in rows or []
        if isinstance(row, (list, tuple)) and any(_normalize_table_cell(cell) for cell in row)
    ]


def _build_table_payload(table_index, rows, *, source, page_number=None, page_table_index=None):
    cleaned_rows = [
        [_normalize_table_cell(cell) for cell in row]
        for row in rows
        if any(_normalize_table_cell(cell) for cell in row)
    ]
    column_count = max((len(row) for row in cleaned_rows), default=0)
    normalized_rows = [row + [""] * (column_count - len(row)) for row in cleaned_rows]
    payload = {
        "table_index": table_index,
        "source": source,
        "row_count": len(normalized_rows),
        "column_count": column_count,
        "headers": normalized_rows[0] if normalized_rows else [],
        "rows": normalized_rows[1:] if len(normalized_rows) > 1 else [],
        "text": _format_table_text(normalized_rows),
    }
    if page_number is not None:
        payload["page_number"] = page_number
    if page_table_index is not None:
        payload["page_table_index"] = page_table_index
    return payload


def _format_table_text(rows):
    return "\n".join(" | ".join(_normalize_table_cell(cell) for cell in row) for row in rows)


def _normalize_table_cell(value):
    return re.sub(r"\s+", " ", str(value or "").strip())


def _parse_text_table_row(line):
    stripped = (line or "").strip()
    if not stripped:
        return None
    if "|" in stripped:
        cells = stripped.strip("|").split("|")
    elif "\t" in stripped:
        cells = stripped.split("\t")
    elif "," in stripped and len(stripped.split(",")) >= 3:
        cells = stripped.split(",")
    else:
        return None

    cells = [_normalize_table_cell(cell) for cell in cells]
    return cells if sum(1 for cell in cells if cell) >= 2 else None


def _get_pdf_table_regions(page):
    regions = []
    for table in _find_pdf_tables(page):
        bbox = getattr(table, "bbox", None)
        if bbox and len(bbox) == 4:
            regions.append(tuple(float(value) for value in bbox))
    for inferred in _infer_pdf_tables_from_words(page):
        bbox = inferred.get("bbox")
        if bbox and len(bbox) == 4:
            regions.append(tuple(float(value) for value in bbox))
    return regions


def _infer_pdf_tables_from_words(page, exclude_bboxes=None):
    exclude_bboxes = exclude_bboxes or []
    try:
        words = page.get_text("words")
    except Exception:
        return []
    if not words:
        return []

    rows = _group_pdf_words_into_rows(words, exclude_bboxes=exclude_bboxes)
    candidate_rows = []
    for row in rows:
        cells = _split_pdf_row_into_cells(row)
        if len(cells) < 3:
            continue
        row_text = " ".join(cell["text"] for cell in cells)
        if len(row_text) > 260:
            continue
        candidate_rows.append(
            {
                "cells": cells,
                "bbox": _merge_bboxes([cell["bbox"] for cell in cells]),
                "numeric_density": _numeric_density(row_text),
            }
        )

    groups = []
    current = []
    for row in candidate_rows:
        if not current:
            current = [row]
            continue
        prev = current[-1]
        vertical_gap = row["bbox"][1] - prev["bbox"][3]
        if vertical_gap <= 18 and _row_columns_are_compatible(prev["cells"], row["cells"]):
            current.append(row)
        else:
            if _looks_like_table_group(current):
                groups.append(current)
            current = [row]
    if _looks_like_table_group(current):
        groups.append(current)

    inferred_tables = []
    for group in groups:
        rows_payload = [[cell["text"] for cell in row["cells"]] for row in group]
        payload = _build_table_payload(len(inferred_tables), rows_payload, source="pdf_inferred")
        payload["bbox"] = _merge_bboxes([row["bbox"] for row in group])
        payload["confidence"] = _estimate_inferred_table_confidence(group)
        inferred_tables.append(payload)
    return inferred_tables


def _group_pdf_words_into_rows(words, *, exclude_bboxes=None):
    rows = []
    for word in sorted(words, key=lambda item: (float(item[1]), float(item[0]))):
        x0, y0, x1, y1 = (float(word[0]), float(word[1]), float(word[2]), float(word[3]))
        if _intersects_any_bbox((x0, y0, x1, y1), exclude_bboxes or []):
            continue
        text = _normalize_table_cell(word[4] if len(word) > 4 else "")
        if not text:
            continue

        placed = False
        center_y = (y0 + y1) / 2
        for row in rows:
            if abs(row["center_y"] - center_y) <= 3.5:
                row["words"].append({"text": text, "bbox": (x0, y0, x1, y1)})
                row["center_y"] = (row["center_y"] + center_y) / 2
                placed = True
                break
        if not placed:
            rows.append({"center_y": center_y, "words": [{"text": text, "bbox": (x0, y0, x1, y1)}]})

    normalized_rows = []
    for row in rows:
        row_words = sorted(row["words"], key=lambda item: item["bbox"][0])
        if row_words:
            normalized_rows.append(row_words)
    return normalized_rows


def _split_pdf_row_into_cells(row_words):
    if not row_words:
        return []

    gaps = []
    for left, right in zip(row_words, row_words[1:]):
        gaps.append(max(0.0, right["bbox"][0] - left["bbox"][2]))
    positive_gaps = [gap for gap in gaps if gap > 0]
    if not positive_gaps:
        return [{"text": " ".join(word["text"] for word in row_words), "bbox": _merge_bboxes([w["bbox"] for w in row_words])}]

    # Large x-axis gaps usually mean columns in borderless/three-line PDF tables.
    # Normal prose extracted by PyMuPDF has word gaps around a few points.
    split_gap = 18.0
    cells = []
    current = [row_words[0]]
    for gap, word in zip(gaps, row_words[1:]):
        if gap >= split_gap:
            cells.append(_build_cell_from_words(current))
            current = [word]
        else:
            current.append(word)
    cells.append(_build_cell_from_words(current))
    return [cell for cell in cells if cell["text"]]


def _build_cell_from_words(words):
    return {
        "text": _normalize_table_cell(" ".join(word["text"] for word in words)),
        "bbox": _merge_bboxes([word["bbox"] for word in words]),
    }


def _row_columns_are_compatible(left_cells, right_cells):
    if abs(len(left_cells) - len(right_cells)) > 1:
        return False
    pairs = zip(left_cells, right_cells)
    aligned = sum(1 for left, right in pairs if abs(left["bbox"][0] - right["bbox"][0]) <= 28)
    return aligned >= min(len(left_cells), len(right_cells), 2)


def _looks_like_table_group(rows):
    if len(rows) < 2:
        return False
    column_rich_rows = [row for row in rows if len(row["cells"]) >= 3]
    if len(column_rich_rows) < 2:
        return False
    numeric_rows = [row for row in rows if row["numeric_density"] >= 0.12]
    repeated_column_count = len({len(row["cells"]) for row in rows}) <= 2
    return repeated_column_count and (len(numeric_rows) >= 1 or len(rows) >= 3)


def _estimate_inferred_table_confidence(rows):
    if not rows:
        return 0.0
    numeric_share = sum(1 for row in rows if row["numeric_density"] >= 0.12) / len(rows)
    row_bonus = min(0.25, max(0, len(rows) - 2) * 0.05)
    return round(min(0.95, 0.55 + numeric_share * 0.25 + row_bonus), 2)


def _numeric_density(text):
    stripped = re.sub(r"\s+", "", text or "")
    if not stripped:
        return 0.0
    numeric_chars = sum(1 for char in stripped if char.isdigit() or char in ".%+-")
    return numeric_chars / len(stripped)


def _merge_bboxes(bboxes):
    boxes = [bbox for bbox in bboxes if bbox and len(bbox) == 4]
    if not boxes:
        return (0.0, 0.0, 0.0, 0.0)
    return (
        min(float(bbox[0]) for bbox in boxes),
        min(float(bbox[1]) for bbox in boxes),
        max(float(bbox[2]) for bbox in boxes),
        max(float(bbox[3]) for bbox in boxes),
    )


def _intersects_any_bbox(block_bbox, table_bboxes):
    return any(_bbox_intersection_ratio(block_bbox, table_bbox) > 0.5 for table_bbox in table_bboxes or [])


def _bbox_intersection_ratio(a, b):
    ax0, ay0, ax1, ay1 = a
    bx0, by0, bx1, by1 = b
    ix0, iy0 = max(ax0, bx0), max(ay0, by0)
    ix1, iy1 = min(ax1, bx1), min(ay1, by1)
    if ix1 <= ix0 or iy1 <= iy0:
        return 0.0
    intersection = (ix1 - ix0) * (iy1 - iy0)
    area = max((ax1 - ax0) * (ay1 - ay0), 1.0)
    return intersection / area


def _docx_paragraph_is_inside_table(paragraph_element):
    current = paragraph_element
    while current is not None:
        if str(getattr(current, "tag", "")).endswith("}tbl"):
            return True
        current = current.getparent()
    return False


def split_text_into_segments(text_content, max_segment_length=1500, overlap_length=100, fallback_segment_length=40000):
    paragraphs = extract_document_paragraphs(text_content)
    segments = []
    
    for paragraph in paragraphs:
        # 尝试整段保留：如果段落本身没有超过最大限制，直接作为完整的一个 segment
        if len(paragraph) <= max_segment_length:
            segments.append(paragraph.strip())
            continue
            
        # 如果段落过长（超过 max_segment_length），则进行句子级的滑动窗口切分
        current_sentences = []
        current_length = 0
        sentence_chunks = _split_paragraph_into_sentences(paragraph)
        
        for chunk in sentence_chunks:
            if len(chunk) > max_segment_length:
                if current_sentences:
                    segments.append(" ".join(current_sentences).strip())
                    current_sentences = []
                    current_length = 0
                
                for hard_chunk in _hard_split_text(chunk, max_segment_length):
                    segments.append(hard_chunk.strip())
                continue

            if current_length + len(chunk) + (1 if current_length else 0) <= max_segment_length:
                current_sentences.append(chunk)
                current_length += len(chunk) + (1 if current_length else 0)
            else:
                if current_sentences:
                    segments.append(" ".join(current_sentences).strip())
                
                # 引入滑动窗口重叠
                overlap_sentences = []
                overlap_len = 0
                for prev_chunk in reversed(current_sentences):
                    if overlap_len + len(prev_chunk) <= overlap_length:
                        overlap_sentences.insert(0, prev_chunk)
                        overlap_len += len(prev_chunk) + 1
                    else:
                        break
                
                current_sentences = overlap_sentences + [chunk]
                current_length = sum(len(s) for s in current_sentences) + len(current_sentences) - 1

        if current_sentences:
            segments.append(" ".join(current_sentences).strip())

    if segments:
        return segments
    if text_content:
        return [sanitize_text_content(text_content[:fallback_segment_length])]
    return [EMPTY_FILE_MESSAGE]


def sanitize_academic_noise(text):
    """
    清洗学术论文中的噪音，包括：图表说明、公式乱码、文内引用。
    这些内容会被替换为统一的占位符，避免大模型误判其为 AI 生成的套话。
    """
    if not text:
        return text

    # 1. 文内引用过滤 (In-text Citation Filtering)
    # 匹配数字上标引用，如 [1], [1, 2], [1-3, 5]
    text = re.sub(r'\[\s*\d+(?:\s*[,，\-]\s*\d+)*\s*\]', '[CITATION]', text)
    # 匹配作者-年份引用，如 (Smith et al., 2023) 或 (王五等, 2021)
    text = re.sub(r'\([A-Za-z\u4e00-\u9fa5\s]+(?:et al\.|等)?[,，]\s*(?:19|20)\d{2}[a-z]?\)', '[CITATION]', text)

    # 4. 直接引语识别 (Direct Quote Identification)
    # 将被双引号包裹的长文本标记出来，大模型看到这个标记会知道这是合规引语
    # 匹配英文字符或中文字符包裹的引号
    text = re.sub(r'("[^"]{30,}")', r'[QUOTE_START] \1 [QUOTE_END]', text)
    text = re.sub(r'(“[^”]{30,}”)', r'[QUOTE_START] \1 [QUOTE_END]', text)

    # 2. 图表说明占位 (Figure & Table Masking)
    # 匹配以 Fig., Figure, Table, 图, 表 开头的整段文本（通常是图表说明）
    # 由于我们是按行/段落处理，如果某一段以这些词开头且长度不算太长，认为是图表说明
    lines = text.split('\n')
    cleaned_lines = []
    for line in lines:
        stripped = line.strip()
        # 匹配 "Fig. 1:", "Figure 2.", "Table III:", "图 1-2" 等开头
        if re.match(r'^(?:Fig\.|Figure|Table|图|表)\s*[\dIVXLCDMivxlcdm\-\.]+\s*[:：\.]?', stripped, flags=re.IGNORECASE):
            # 如果段落长度小于 300 字符，大概率是图表说明，直接替换
            if len(stripped) < 300:
                if re.match(r'^(?:Table|表)', stripped, flags=re.IGNORECASE):
                    cleaned_lines.append('[TABLE_CAPTION_REMOVED]')
                else:
                    cleaned_lines.append('[FIGURE_CAPTION_REMOVED]')
                continue
                
        # 3. 公式占位 (Formula Masking)
        # 如果一行文本中包含较多等号或常见数学符号，且不包含太多正常单词，可能是 PDF 提取失败的独立公式行
        # 例如: "E = m c^2 (1)"
        if re.match(r'^.*?(?:=|\+|-|\*|/|\\alpha|\\beta|\\sum|\\int).*?(?:\(\d+\))?$', stripped):
            # 简单的启发式：如果字母数量很少，且有数学符号，或者以 (数字) 结尾
            letters = sum(c.isalpha() for c in stripped)
            if len(stripped) > 5 and letters / len(stripped) < 0.4 and ('=' in stripped or '+' in stripped):
                cleaned_lines.append('[FORMULA_REMOVED]')
                continue

        cleaned_lines.append(line)

    return '\n'.join(cleaned_lines)


def extract_document_paragraphs(text_content):
    raw_text = sanitize_text_content(text_content or "")
    if not raw_text:
        return []

    normalized_text = re.sub(r"\r\n?", "\n", raw_text)
    paragraphs = []
    current_parts = []

    for raw_line in normalized_text.split("\n"):
        cleaned_line = re.sub(r"\s+", " ", raw_line.strip())
        if not cleaned_line:
            if current_parts:
                paragraphs.append(" ".join(current_parts).strip())
                current_parts = []
            continue

        if not current_parts:
            current_parts.append(cleaned_line)
            continue

        previous_line = current_parts[-1]
        if _should_start_new_paragraph(previous_line, cleaned_line):
            paragraphs.append(" ".join(current_parts).strip())
            current_parts = [cleaned_line]
            continue

        current_parts.append(cleaned_line)

    if current_parts:
        paragraphs.append(" ".join(current_parts).strip())

    return [paragraph for paragraph in paragraphs if paragraph]


def parse_document_sections(text_content):
    """将文档精细化拆分为：摘要、正文、致谢、参考文献"""
    paragraphs = extract_document_paragraphs(text_content)
    sections = {
        "abstract": [],
        "body": [],
        "acknowledgements": [],
        "references": []
    }
    
    current_section = "body"
    
    for paragraph in paragraphs:
        p_lower = paragraph.lower().strip()
        # 移除前导的数字和特殊符号，如 "1. " 或 "10 " 或 "IV."
        p_clean = re.sub(r'^[\dIVXLCDMivxlcdm]+[\.\s]*', '', p_lower).strip()
        
        # 匹配标题的严格正则：要求整个段落非常短（通常是标题的特征）
        # 并且完全匹配这些关键字（忽略大小写和前导数字后）
        is_heading = len(p_clean) < 50
        
        if is_heading and p_clean in {"abstract", "摘要"}:
            current_section = "abstract"
            continue
        elif is_heading and p_clean in {"introduction", "引言", "导言", "绪论"}:
            current_section = "body"
            sections[current_section].append(paragraph)
            continue
        elif is_heading and p_clean in {"acknowledgements", "acknowledgments", "致谢", "致谢辞"}:
            current_section = "acknowledgements"
            continue
        elif is_heading and p_clean in {"references", "bibliography", "参考文献", "参考书目"}:
            current_section = "references"
            continue
            
        sections[current_section].append(paragraph)
        
    return {k: "\n".join(v) for k, v in sections.items()}


def extract_document_references(text_content):
    sections = parse_document_sections(text_content)
    if sections["references"].strip():
        raw_refs = [p.strip() for p in sections["references"].split("\n") if p.strip()]
        merged_refs = []
        for ref in raw_refs:
            # 判断是否像是一个新参考文献的开头，比如 "[1]", "1.", "(1)"
            if re.match(r'^(\[\d+\]|\(\d+\)|\d+\.)', ref) or not merged_refs:
                merged_refs.append(ref)
            else:
                merged_refs[-1] += " " + ref
        return merged_refs

    # Fallback heuristic
    paragraphs = extract_document_paragraphs(text_content)
    return [
        paragraph
        for paragraph in paragraphs
        if paragraph.startswith("[") or paragraph[:2].isdigit() or "doi" in paragraph.lower()
    ]





def _split_paragraph_into_sentences(paragraph):
    text = sanitize_text_content(paragraph or "")
    if not text:
        return []

    parts = re.split(r"(?<=[。！？!?；;\.])\s+", text)
    refined = []
    for part in parts:
        cleaned = part.strip()
        if cleaned:
            refined.append(cleaned)
    return refined or [text]


def _hard_split_text(text, max_len):
    stripped = (text or "").strip()
    if not stripped:
        return []
        
    chunks = []
    while len(stripped) > max_len:
        split_at = stripped.rfind(" ", 0, max_len)
        if split_at == -1:
            split_at = max_len
        chunks.append(stripped[:split_at].strip())
        stripped = stripped[split_at:].strip()
        
    if stripped:
        chunks.append(stripped)
    return chunks


def _should_start_new_paragraph(previous_line, current_line):
    if not previous_line.strip():
        return True
    if _looks_like_reference_item(current_line):
        return True
    if _looks_like_heading(current_line):
        return True
    if _looks_like_list_item(current_line) and not _looks_like_sentence_continuation(previous_line, current_line):
        return True
    if re.search(r'[。！？!?；;:：.]["”’)\]]*$', previous_line) and _looks_like_sentence_start(current_line):
        return True
    return False


def _looks_like_reference_item(line):
    return bool(re.match(r'^(\[\d+\]|\(\d+\)|\d+\.)\s*\S', line))


def _looks_like_list_item(line):
    return bool(
        re.match(
            r'^(?:[-*•]|(?:\d+|[A-Za-z]|[IVXLCDMivxlcdm]+)[\.\)])\s+\S',
            line,
        )
    )


def _looks_like_heading(line):
    stripped = line.strip()
    normalized = stripped.lower()
    normalized = re.sub(r'^[\dIVXLCDMivxlcdm]+[\.\s]*', '', normalized).strip()
    if normalized in {
        "abstract",
        "introduction",
        "background",
        "method",
        "methods",
        "experiment",
        "experiments",
        "results",
        "discussion",
        "conclusion",
        "acknowledgements",
        "acknowledgments",
        "references",
        "bibliography",
        "摘要",
        "引言",
        "导言",
        "绪论",
        "方法",
        "实验",
        "结果",
        "结论",
        "致谢",
        "参考文献",
        "参考书目",
    }:
        return True
    if len(stripped) > 60:
        return False
    if re.search(r'[。！？!?；;:：,，]', stripped):
        return False
    return bool(re.match(r'^[A-Z][A-Z0-9\s\-]{2,}$', stripped))


def _looks_like_sentence_start(line):
    stripped = line.lstrip()
    if not stripped:
        return False
    first_char = stripped[0]
    if "\u4e00" <= first_char <= "\u9fff":
        return True
    return first_char.isupper() or first_char.isdigit() or first_char in {'"', "'", "“", "‘", "(", "["}


def _looks_like_sentence_continuation(previous_line, current_line):
    prev = previous_line.rstrip()
    cur = current_line.lstrip()
    if prev.endswith("-"):
        return True
    if not cur:
        return False
    first_char = cur[0]
    if first_char.islower():
        return True
    return first_char in {",", ".", ";", ":", ")", "]"}
